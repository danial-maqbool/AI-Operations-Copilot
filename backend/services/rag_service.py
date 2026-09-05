import re
import shutil
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sqlalchemy.orm import Session
import pypdf
import docx

from backend.config import settings
from backend.models import Document, DocumentChunk, Workspace
from backend.schemas.document import SearchResult

class RAGService:
    @classmethod
    def extract_text_from_pdf(cls, file_path: Path) -> List[Dict[str, Any]]:
        pages_data = []
        reader = pypdf.PdfReader(str(file_path))
        for idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages_data.append({
                "page_number": idx + 1,
                "text": text.strip()
            })
        return pages_data

    @classmethod
    def extract_text_from_docx(cls, file_path: Path) -> List[Dict[str, Any]]:
        doc = docx.Document(str(file_path))
        pages_data = []
        curr_section = "General"
        curr_text = []
        page_num = 1
        word_count = 0

        for p in doc.paragraphs:
            t = p.text.strip()
            if not t:
                continue
            if p.style and "Heading" in p.style.name:
                curr_section = t
            curr_text.append(t)
            word_count += len(t.split())
            if word_count > 400:
                pages_data.append({
                    "page_number": page_num,
                    "section": curr_section,
                    "text": "\n".join(curr_text)
                })
                page_num += 1
                curr_text = []
                word_count = 0

        if curr_text:
            pages_data.append({
                "page_number": page_num,
                "section": curr_section,
                "text": "\n".join(curr_text)
            })
        return pages_data

    @classmethod
    def extract_text_from_text(cls, file_path: Path) -> List[Dict[str, Any]]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        # Split into sections based on Markdown headers
        sections = re.split(r'\n(?=#{1,3}\s+)', content)
        pages_data = []
        page_num = 1

        for s in sections:
            s_clean = s.strip()
            if not s_clean:
                continue
            # Check for header
            header_match = re.match(r'^#{1,3}\s+(.+)$', s_clean, re.MULTILINE)
            sec_name = header_match.group(1).strip() if header_match else "General"
            
            # Estimate pages if section is long
            words = s_clean.split()
            chunk_size = 400
            for i in range(0, max(len(words), 1), chunk_size):
                sub_text = " ".join(words[i:i+chunk_size])
                pages_data.append({
                    "page_number": page_num,
                    "section": sec_name,
                    "text": sub_text
                })
                page_num += 1

        return pages_data

    @classmethod
    def chunk_page_data(cls, pages_data: List[Dict[str, Any]], chunk_words: int = 250, overlap_words: int = 40) -> List[Dict[str, Any]]:
        chunks = []
        chunk_idx = 0

        for p_info in pages_data:
            page_no = p_info["page_number"]
            sec_title = p_info.get("section", "Standard Operating Procedure")
            text = p_info["text"]
            words = text.split()

            if not words:
                continue

            # If small enough, single chunk
            if len(words) <= chunk_words:
                chunks.append({
                    "chunk_index": chunk_idx,
                    "page_number": page_no,
                    "section_title": sec_title,
                    "content": text,
                    "token_count": len(words)
                })
                chunk_idx += 1
            else:
                start = 0
                while start < len(words):
                    end = min(start + chunk_words, len(words))
                    chunk_text = " ".join(words[start:end])
                    chunks.append({
                        "chunk_index": chunk_idx,
                        "page_number": page_no,
                        "section_title": sec_title,
                        "content": chunk_text,
                        "token_count": len(words[start:end])
                    })
                    chunk_idx += 1
                    start += (chunk_words - overlap_words)

        return chunks

    @classmethod
    def ingest_document(cls, file_path: Path, filename: str, workspace_id: Optional[str], db: Session) -> Document:
        ws = db.query(Workspace).filter(Workspace.id == workspace_id).first() if workspace_id else db.query(Workspace).first()
        if not ws:
            ws = Workspace(name="Default Workspace")
            db.add(ws)
            db.commit()

        ext = file_path.suffix.lower()
        if ext == ".pdf":
            pages_data = cls.extract_text_from_pdf(file_path)
            doc_type = "pdf"
        elif ext in [".docx", ".doc"]:
            pages_data = cls.extract_text_from_docx(file_path)
            doc_type = "docx"
        elif ext in [".txt", ".md"]:
            pages_data = cls.extract_text_from_text(file_path)
            doc_type = "txt" if ext == ".txt" else "md"
        else:
            raise ValueError(f"Unsupported document format: {ext}")

        chunks_data = cls.chunk_page_data(pages_data)

        # Check existing document
        existing = db.query(Document).filter(
            Document.workspace_id == ws.id,
            Document.filename == filename
        ).first()

        if existing:
            doc = existing
            doc.total_pages = len(pages_data)
            doc.total_chunks = len(chunks_data)
            doc.status = "INDEXED"
            # Delete old chunks
            db.query(DocumentChunk).filter(DocumentChunk.document_id == doc.id).delete()
        else:
            doc = Document(
                workspace_id=ws.id,
                filename=filename,
                file_type=doc_type,
                file_path=str(file_path),
                total_pages=len(pages_data),
                total_chunks=len(chunks_data),
                status="INDEXED"
            )
            db.add(doc)
            db.flush()

        for c in chunks_data:
            chunk_rec = DocumentChunk(
                document_id=doc.id,
                chunk_index=c["chunk_index"],
                page_number=c["page_number"],
                section_title=c["section_title"],
                content=c["content"],
                token_count=c["token_count"]
            )
            db.add(chunk_rec)

        db.commit()
        db.refresh(doc)
        return doc

    @classmethod
    def search(cls, query: str, workspace_id: Optional[str], db: Session, top_k: int = 4) -> List[SearchResult]:
        ws = db.query(Workspace).filter(Workspace.id == workspace_id).first() if workspace_id else db.query(Workspace).first()
        if not ws:
            return []

        chunks = (
            db.query(DocumentChunk)
            .join(Document)
            .filter(Document.workspace_id == ws.id)
            .all()
        )

        if not chunks:
            return []

        corpus = [c.content for c in chunks]
        
        # High-performance TF-IDF vector search
        try:
            vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
            tfidf_matrix = vectorizer.fit_transform(corpus)
            query_vec = vectorizer.transform([query])
            sim_scores = cosine_similarity(query_vec, tfidf_matrix).flatten()
        except Exception:
            return []

        # Get top indices
        top_indices = np.argsort(sim_scores)[::-1][:top_k]
        results = []

        for idx in top_indices:
            score = float(sim_scores[idx])
            if score > 0.05:  # Relevance threshold
                chunk = chunks[idx]
                doc = chunk.document
                citation_str = (
                    f"Source: {doc.filename}\n"
                    f"Page: {chunk.page_number}\n"
                    f"Section: {chunk.section_title or 'General'}"
                )
                results.append(SearchResult(
                    document_id=doc.id,
                    filename=doc.filename,
                    page_number=chunk.page_number,
                    section_title=chunk.section_title,
                    content=chunk.content,
                    score=round(score, 3),
                    citation=citation_str
                ))

        return results
